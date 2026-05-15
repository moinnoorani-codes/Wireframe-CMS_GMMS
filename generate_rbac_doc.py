from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page Setup (Portrait) ───
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(1)

ACTION_LABELS = {
    'view': 'View',
    'create': 'Create',
    'update': 'Update',
    'delete': 'Delete',
    'export': 'Export/Download',
    'print': 'Print',
}

ACTION_ORDER = ['view', 'create', 'update', 'delete', 'export', 'print']

# ─── Helper functions ───
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_header_row(table, headers, font_size=8):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'
        set_cell_shading(cell, '2F5496')

def add_alt_row(table, values, font_size=8, bold_first=False):
    row = table.add_row()
    row_idx = len(table.rows) - 1
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
                if bold_first and i == 0:
                    run.bold = True
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'F2F2F2')
    return row

def add_section_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        run.font.name = 'Calibri'

def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

# ═══════════════════════════════════════════════
#  DATA: SERVICE MODULES WITH CANONICAL ACTIONS
# ═══════════════════════════════════════════════

services = [
    {
        'code': 'S01', 'name': 'Dashboard',
        'actions': {
            'view': 'View the main dashboard, KPIs, charts, activity feed, and alerts',
        }
    },
    {
        'code': 'S02', 'name': 'Products & SKU',
        'actions': {
            'view': 'Browse the SKU catalog, search, filter, view details',
            'create': 'Create new SKU entries (both GMMS and External)',
            'update': 'Modify existing SKU details, pricing, stock, images',
            'delete': 'Remove SKU entries from the system',
            'export': 'Export SKU/product data',
            'print': 'Print barcode/label sheets for SKUs',
        }
    },
    {
        'code': 'S03', 'name': 'Inventory / Stock',
        'actions': {
            'view': 'Browse live inventory with location breakdown, view stock alerts and ageing report',
            'update': 'Adjust stock quantities (add/reduce)',
            'export': 'Export inventory data',
        }
    },
    {
        'code': 'S04', 'name': 'Orders (Retail & Wholesale)',
        'actions': {
            'view': 'Browse order list, view order details, filter by status',
            'create': 'Create new retail and wholesale orders',
            'update': 'Edit orders, approve/reject wholesale orders, record payments',
            'delete': 'Cancel or delete orders',
            'export': 'Export order data',
            'print': 'Print challan documents for orders',
        }
    },
    {
        'code': 'S05', 'name': 'Dispatch & Logistics',
        'actions': {
            'view': 'View LR management console and LR documents',
            'create': 'Upload LR documents (photo/PDF)',
            'update': 'Pick items, confirm dispatch, resend SMS notifications',
            'print': 'Print challan during dispatch',
        }
    },
    {
        'code': 'S06', 'name': 'CCTV',
        'actions': {
            'view': 'Browse and watch recorded CCTV footage (footage library access restricted — Godown Staff can record but cannot browse footage)',
            'create': 'Record dispatch video via CCTV cameras',
            'export': 'Download video clips from the footage library',
        }
    },
    {
        'code': 'S07', 'name': 'Payments & Finance',
        'actions': {
            'view': 'View payment records and outstanding balances',
            'create': 'Record new payments',
            'update': 'Perform daily reconciliation, close day-end operations',
            'export': 'Export payment and reconciliation data',
        }
    },
    {
        'code': 'S08', 'name': 'Reports & Analytics',
        'actions': {
            'view': 'Access reports hub, view sales, ageing, top designs, customer history reports',
            'export': 'Export reports as PDF or CSV',
        }
    },
    {
        'code': 'S09', 'name': 'SMS Communication',
        'actions': {
            'view': 'View SMS delivery log and status',
            'create': 'Manually trigger SMS sending',
            'update': 'Create, edit, activate/deactivate SMS templates',
        }
    },
    {
        'code': 'S10', 'name': 'CRM / Customer Management',
        'actions': {
            'view': 'Browse customer master list with details',
            'create': 'Register new wholesale/retail customers',
            'update': 'Modify customer details, credit limits, discounts',
            'delete': 'Remove customer records',
            'export': 'Export customer data',
        }
    },
    {
        'code': 'S11', 'name': 'Mfg — Challans',
        'actions': {
            'view': 'Browse challan list, view tracking timeline and progress',
            'create': 'Create new production challans',
            'update': 'Modify existing challans, close challans',
            'delete': 'Delete challans',
            'print': 'Print challan documents',
        }
    },
    {
        'code': 'S12', 'name': 'Mfg — Production',
        'actions': {
            'view': 'View production status, counts, and contractor payment checks',
            'update': 'Update piece counts, process SKU outward',
            'export': 'Export production data',
        }
    },
    {
        'code': 'S13', 'name': 'Mfg — Contractors',
        'actions': {
            'view': 'Browse contractor list and details',
            'create': 'Register new contractors',
            'update': 'Modify contractor details',
            'delete': 'Remove contractor records',
        }
    },
    {
        'code': 'S14', 'name': 'Mfg — Fabric / Mill',
        'actions': {
            'view': 'View fabric roll inventory and mill data',
            'create': 'Add new fabric rolls to inventory',
            'update': 'Edit fabric roll details',
        }
    },
    {
        'code': 'S15', 'name': 'Mfg — RF / Returns',
        'actions': {
            'view': 'View return fabric (RF) entries',
            'create': 'Create new RF entries for returns',
            'update': 'Modify existing RF entries',
        }
    },
    {
        'code': 'S16', 'name': 'Mfg — Masters (Design / Job / Color)',
        'actions': {
            'view': 'View design masters, job work types, color configurations',
            'update': 'Upload/manage DST design files, configure job types/rates, manage colors',
            'export': 'Export design master data and configurations',
            'print': 'Print design specification sheets',
        }
    },
    {
        'code': 'S17', 'name': 'Mfg — Costing',
        'actions': {
            'view': 'View design BOM and cost breakdown',
            'update': 'Modify cost data',
            'export': 'Export costing data',
        }
    },
    {
        'code': 'S18', 'name': 'Mfg — Notifications',
        'actions': {
            'view': 'View manufacturing notification center',
        }
    },
    {
        'code': 'S19', 'name': 'Admin & Settings',
        'actions': {
            'view': 'View audit trail and activity log',
            'update': 'Manage users, roles, permissions, and system settings',
            'export': 'Export audit trail data',
        }
    },
    {
        'code': 'S20', 'name': 'Mobile App (Godown)',
        'actions': {
            'view': 'Access mobile Godown app, barcode scan, view orders',
            'create': 'Record dispatch video via mobile CCTV',
            'update': 'Process dispatch, picking, LR upload, manage customers on mobile',
        }
    },
    {
        'code': 'S21', 'name': 'Contractor App',
        'actions': {
            'view': 'View assigned challans, payment ledger and history',
            'update': 'Accept/reject challans, confirm pieces, edit profile',
        }
    },
]

# ═══════════════════════════════════════════════
#  DATA: PREDEFINED ROLES
# ═══════════════════════════════════════════════

roles = {
    'Super Admin': {
        'desc': 'Full unrestricted access to all modules and features including costing, admin settings, and sensitive operations. Typically assigned to the business owner or system administrator.',
        'perms': 'ALL'
    },
    'Manager': {
        'desc': 'Operational control across Sales ERP and Manufacturing ERP. Can manage most day-to-day operations but restricted from sensitive admin settings.',
        'perms': [
            'dashboard:view',
            'products:view', 'products:create', 'products:update', 'products:delete', 'products:export', 'products:print',
            'inventory:view', 'inventory:update', 'inventory:export',
            'orders:view', 'orders:create', 'orders:update', 'orders:delete', 'orders:export', 'orders:print',
            'dispatch:view', 'dispatch:create', 'dispatch:update', 'dispatch:print',
            'cctv:view', 'cctv:create', 'cctv:export',
            'payments:view', 'payments:create', 'payments:update', 'payments:export',
            'reports:view', 'reports:export',
            'sms:view', 'sms:create', 'sms:update',
            'crm:view', 'crm:create', 'crm:update', 'crm:delete', 'crm:export',
            'mfg_challan:view', 'mfg_challan:create', 'mfg_challan:update', 'mfg_challan:delete', 'mfg_challan:print',
            'mfg_production:view', 'mfg_production:update', 'mfg_production:export',
            'mfg_contractor:view', 'mfg_contractor:create', 'mfg_contractor:update', 'mfg_contractor:delete',
            'mfg_fabric:view', 'mfg_fabric:create', 'mfg_fabric:update',
            'mfg_rf:view', 'mfg_rf:create', 'mfg_rf:update',
            'mfg_masters:view', 'mfg_masters:update', 'mfg_masters:export', 'mfg_masters:print',
            'mfg_notifications:view',
            'mobile:view', 'mobile:create', 'mobile:update',
            'contractor:view',
            'admin:view',
        ]
    },
    'Salesman': {
        'desc': 'Frontline sales role. Can create retail and wholesale orders, manage customers, view products, record payments. Cannot approve orders, manage inventory, or access manufacturing.',
        'perms': [
            'dashboard:view',
            'products:view',
            'orders:view', 'orders:create', 'orders:update', 'orders:export', 'orders:print',
            'crm:view', 'crm:create', 'crm:update',
            'payments:view', 'payments:create',
            'reports:view', 'reports:export',
            'sms:view',
            'mobile:view',
        ]
    },
    'Office Staff': {
        'desc': 'General office operations. Handles order processing, payment recording, dispatch tracking, and basic customer management.',
        'perms': [
            'dashboard:view',
            'products:view',
            'inventory:view',
            'orders:view', 'orders:create', 'orders:update', 'orders:export', 'orders:print',
            'dispatch:view',
            'payments:view', 'payments:create',
            'reports:view', 'reports:export',
            'crm:view', 'crm:update',
            'mobile:view',
        ]
    },
    'Godown Staff': {
        'desc': 'Warehouse / godown operations. Handles inventory, order picking, dispatch confirmation, LR upload, and mobile scanning. No sales or financial access.',
        'perms': [
            'dashboard:view',
            'products:view', 'products:print',
            'inventory:view', 'inventory:update', 'inventory:export',
            'orders:view', 'orders:print',
            'dispatch:view', 'dispatch:update', 'dispatch:print',
            'cctv:view', 'cctv:create',
            'mobile:view', 'mobile:create', 'mobile:update',
        ]
    },
    'Accounts': {
        'desc': 'GMMS finance and accounts role. Manages payments, reconciliation, financial reports, and SMS communication for Manufacturing ERP. No order creation or production module access.',
        'perms': [
            'dashboard:view',
            'products:view',
            'orders:view', 'orders:export',
            'payments:view', 'payments:create', 'payments:update', 'payments:export',
            'reports:view', 'reports:export',
            'sms:view', 'sms:create',
            'crm:view',
        ]
    },
    'Production Manager': {
        'desc': 'Full manufacturing ERP access. Manages challans, contractors, production flow, fabric, RF returns, and masters. No Sales ERP access except read-only viewing.',
        'perms': [
            'dashboard:view',
            'products:view',
            'inventory:view',
            'orders:view',
            'mfg_challan:view', 'mfg_challan:create', 'mfg_challan:update', 'mfg_challan:delete', 'mfg_challan:print',
            'mfg_production:view', 'mfg_production:update', 'mfg_production:export',
            'mfg_contractor:view', 'mfg_contractor:create', 'mfg_contractor:update', 'mfg_contractor:delete',
            'mfg_fabric:view', 'mfg_fabric:create', 'mfg_fabric:update',
            'mfg_rf:view', 'mfg_rf:create', 'mfg_rf:update',
            'mfg_masters:view', 'mfg_masters:update', 'mfg_masters:export', 'mfg_masters:print',
            'mfg_notifications:view',
            'reports:view',
            'mobile:view',
        ]
    },
    'Production Staff': {
        'desc': 'Manufacturing floor staff. Can view challans, update piece counts, process SKU outward, and view contractor info. No create/delete/edit on masters.',
        'perms': [
            'dashboard:view',
            'products:view',
            'inventory:view',
            'mfg_challan:view',
            'mfg_production:view', 'mfg_production:update',
            'mfg_contractor:view',
            'mfg_fabric:view',
            'mfg_rf:view',
            'mfg_notifications:view',
            'mobile:view',
        ]
    },
    'Contractor': {
        'desc': 'External contractor mobile app access. Can view assigned challans, accept/reject, submit pieces, view payments, and edit own profile. No access to Sales ERP or main system.',
        'perms': [
            'contractor:view', 'contractor:update',
        ]
    },
    'Read-Only': {
        'desc': 'View-only access across all modules. Can browse any screen, view data, and export reports. No create, update, delete, or print operations.',
        'perms': [
            'dashboard:view',
            'products:view',
            'inventory:view',
            'orders:view',
            'dispatch:view',
            'payments:view',
            'reports:view', 'reports:export',
            'sms:view',
            'crm:view',
            'mfg_challan:view',
            'mfg_production:view',
            'mfg_contractor:view',
            'mfg_fabric:view',
            'mfg_rf:view',
            'mfg_notifications:view',
            'mobile:view',
        ]
    },
}

# Build lookup: module prefix -> module data
module_lookup = {}
for svc in services:
    key = svc['name'].lower().split(' — ')[0].split(' / ')[0].split(' (')[0].replace(' & ', '_').replace(' ', '_').replace('—', '').strip('_')
    # Better: use a manual short prefix
    module_lookup[svc['code']] = svc

# Map short module prefixes to code
PREFIX_TO_CODE = {}
for svc in services:
    parts = svc['name'].lower().replace(' — ', '_').replace(' & ', '_').replace(' / ', '_').replace(' ', '_').replace('(', '').replace(')', '').split('_')
    # First meaningful part
    prefix = parts[0]
    if prefix in ('mfg',):
        PREFIX_TO_CODE[f'mfg_{parts[1]}'] = svc['code']
    elif prefix == 'mobile':
        PREFIX_TO_CODE['mobile'] = svc['code']
    elif prefix == 'contractor':
        PREFIX_TO_CODE['contractor'] = svc['code']
    elif prefix == 'admin':
        PREFIX_TO_CODE['admin'] = svc['code']
    else:
        PREFIX_TO_CODE[prefix] = svc['code']

# Manual mapping since automatic extraction is fragile
PERM_PREFIX_MAP = {
    'dashboard': 'S01',
    'products': 'S02',
    'inventory': 'S03',
    'orders': 'S04',
    'dispatch': 'S05',
    'cctv': 'S06',
    'payments': 'S07',
    'reports': 'S08',
    'sms': 'S09',
    'crm': 'S10',
    'mfg_challan': 'S11',
    'mfg_production': 'S12',
    'mfg_contractor': 'S13',
    'mfg_fabric': 'S14',
    'mfg_rf': 'S15',
    'mfg_masters': 'S16',
    'mfg_costing': 'S17',
    'mfg_notifications': 'S18',
    'admin': 'S19',
    'mobile': 'S20',
    'contractor': 'S21',
}

def get_service_for_perm(perm_key):
    prefix = perm_key.split(':')[0]
    code = PERM_PREFIX_MAP.get(prefix)
    if code:
        for svc in services:
            if svc['code'] == code:
                return svc
    return None

def role_has_perm(role_name, perm_key):
    if role_name == 'Super Admin':
        return True
    role_data = roles[role_name]
    return perm_key in role_data['perms']

# ═══════════════════════════════════════════════
#  BUILD THE DOCUMENT
# ═══════════════════════════════════════════════

# ─── COVER PAGE ───
for _ in range(8):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HOORTEX — Role Based Access Control')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Roles & Permissions Matrix')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
lines = [
    'Document Type:  BA Specification Deliverable',
    'System:         HoorTex CMS + GMMS (Sales ERP + Manufacturing ERP)',
    'Action Model:   6 canonical actions — View · Create · Update · Delete · Export/Download · Print',
    'Platforms:      Web (Sales ERP) · Web (Manufacturing ERP) · Mobile Godown App · Contractor Mobile App',
    'Date:           May 2026',
    'Status:         Draft v2.0',
]
for line in lines:
    run = meta.add_run(line + '\n')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.font.name = 'Calibri'

doc.add_page_break()

# ─── TABLE OF CONTENTS ───
add_section_heading('Table of Contents', 1)
toc = [
    '1.  RBAC Architecture & Key Concepts',
    '2.  Module Action Reference',
    '    2.1  Dashboard',
    '    2.2  Products & SKU',
    '    2.3  Inventory / Stock',
    '    2.4  Orders (Retail & Wholesale)',
    '    2.5  Dispatch & Logistics',
    '    2.6  CCTV',
    '    2.7  Payments & Finance',
    '    2.8  Reports & Analytics',
    '    2.9  SMS Communication',
    '    2.10 CRM / Customer Management',
    '    2.11 Mfg — Challans',
    '    2.12 Mfg — Production',
    '    2.13 Mfg — Contractors',
    '    2.14 Mfg — Fabric / Mill',
    '    2.15 Mfg — RF / Returns',
    '    2.16 Mfg — Masters (Design / Job / Color)',
    '    2.17 Mfg — Costing',
    '    2.18 Mfg — Notifications',
    '    2.19 Admin & Settings',
    '    2.20 Mobile App (Godown)',
    '    2.21 Contractor App',
    '3.  Predefined Roles with Action Bundles',
    '4.  Role vs. Module-Action Matrix',
    '5.  Revision History & Notes',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════
#  SECTION 1: RBAC ARCHITECTURE
# ═══════════════════════════════════════════════
add_section_heading('1.  RBAC Architecture & Key Concepts', 1)

add_section_heading('1.1  Canonical Action Model', 2)
add_body('Every granular permission in the system is defined using exactly one of six canonical actions:')

action_tbl = doc.add_table(rows=1, cols=2)
action_tbl.style = 'Table Grid'
action_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in action_tbl.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(13.0)
add_header_row(action_tbl, ['Action', 'Scope of Access'], font_size=9)

action_descriptions = [
    ('View', 'Browse screens, view data tables and details, search, filter, navigate'),
    ('Create', 'Add new records, register entries, upload documents, record new transactions'),
    ('Update', 'Edit existing records, modify details, approve/reject, process operations, close transactions'),
    ('Delete', 'Remove records, cancel/delete entries from the system'),
    ('Export/Download', 'Export data to file formats, download documents, footage, and attachments'),
    ('Print', 'Print documents, labels, challans, and specification sheets'),
]
for action, desc in action_descriptions:
    add_alt_row(action_tbl, [action, desc], font_size=9, bold_first=True)

doc.add_paragraph()

add_section_heading('1.2  Permission Key Convention', 2)
add_body('Every permission follows the format:')
p = doc.add_paragraph()
run = p.add_run('{module}:{action}')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

p2 = doc.add_paragraph()
run = p2.add_run('Examples:  products:view   ·   orders:create   ·   dispatch:update   ·   mfg_challan:delete')
run.font.size = Pt(9.5)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

add_section_heading('1.3  How Roles Work', 2)
items = [
    'Each Predefined Role is a bundle of module:action permission keys.',
    'When a user is assigned a role, they automatically inherit ALL permissions bundled under that role.',
    'Permissions are additive. If a custom role is created, individual permission keys can be toggled on/off.',
    'A user can be assigned exactly ONE primary role (no role stacking).',
    'The Super Admin role has ALL permissions — it bypasses individual key checks entirely.',
]
for item in items:
    add_bullet(item)

add_section_heading('1.4  Total Count', 2)
total_perms = sum(len(svc['actions']) for svc in services)
total_role_count = len(roles)
p = doc.add_paragraph()
run = p.add_run(f'{total_perms} permission keys  ·  21 service modules  ·  {total_role_count} predefined roles  ·  6 canonical actions')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_page_break()

# ═══════════════════════════════════════════════
#  SECTION 2: MODULE ACTION REFERENCE
# ═══════════════════════════════════════════════
add_section_heading('2.  Module Action Reference', 1)
add_body('Each module supports a subset of the 6 canonical actions. The table below shows which actions are available for each module and what they enable.')

for svc in services:
    add_section_heading(f'{svc["code"]}  {svc["name"]}', 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.5)
    add_header_row(tbl, ['Action', 'Description'], font_size=8)
    for action in ACTION_ORDER:
        if action in svc['actions']:
            add_alt_row(tbl, [ACTION_LABELS[action], svc['actions'][action]], font_size=8, bold_first=True)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════
#  SECTION 3: PREDEFINED ROLES
# ═══════════════════════════════════════════════
add_section_heading('3.  Predefined Roles with Action Bundles', 1)
add_body(f'The following {total_role_count} predefined roles are built into the system. Each role grants a specific set of module:action permissions.')

role_order = ['Super Admin', 'Manager', 'Salesman', 'Office Staff', 'Godown Staff',
              'Accounts', 'Production Manager', 'Production Staff', 'Contractor', 'Read-Only']

for role_name in role_order:
    rd = roles[role_name]
    add_section_heading(f'{role_name}', 2)
    p = doc.add_paragraph()
    run = p.add_run(rd['desc'])
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.italic = True

    if rd['perms'] == 'ALL':
        p2 = doc.add_paragraph()
        run = p2.add_run('Grants ALL permission keys across all 21 service modules.')
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    else:
        from collections import OrderedDict
        grouped = OrderedDict()
        for pk in rd['perms']:
            svc_code = PERM_PREFIX_MAP.get(pk.split(':')[0], '??')
            action = pk.split(':')[1]
            if svc_code not in grouped:
                grouped[svc_code] = []
            grouped[svc_code].append(action)

        p_count = len(rd['perms'])
        svc_count = len(grouped)
        p = doc.add_paragraph()
        run = p.add_run(f'Total: {p_count} permissions across {svc_count} service modules')
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'

        for svc_code, actions in grouped.items():
            svc_name = svc_code
            for s in services:
                if s['code'] == svc_code:
                    svc_name = s['name']
                    break
            p = doc.add_paragraph()
            run = p.add_run(f'{svc_code} — {svc_name}:  ')
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Calibri'
            labels = [ACTION_LABELS.get(a, a) for a in actions]
            run = p.add_run(', '.join(labels))
            run.font.size = Pt(8)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)

    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════
#  SECTION 4: ROLE vs MODULE-ACTION MATRIX
# ═══════════════════════════════════════════════
add_section_heading('4.  Role vs. Module-Action Matrix', 1)
add_body('The following matrix shows which module:action permissions each predefined role includes. '
         'A filled cell (✓) indicates the role has that action for the given module. '
         'This is the master reference for both development and UI implementation.')

abbrev = {
    'Super Admin': 'Super\nAdmin',
    'Manager': 'Manager',
    'Salesman': 'Salesman',
    'Office Staff': 'Office\nStaff',
    'Godown Staff': 'Godown\nStaff',
    'Accounts': 'Accounts',
    'Production Manager': 'Prod\nManager',
    'Production Staff': 'Prod\nStaff',
    'Contractor': 'Contractor',
    'Read-Only': 'Read\nOnly',
}

role_names_short = list(roles.keys())

for svc in services:
    available_actions = [a for a in ACTION_ORDER if a in svc['actions']]
    if not available_actions:
        continue

    add_section_heading(f'{svc["code"]} — {svc["name"]}', 3)
    headers = ['Action'] + [abbrev[r] for r in role_names_short]
    tbl = doc.add_table(rows=1, cols=1 + len(role_names_short))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tbl, headers, font_size=7)

    for action in available_actions:
        perm_key = f'{list(PERM_PREFIX_MAP.keys())[list(PERM_PREFIX_MAP.values()).index(svc["code"])]}:{action}'
        # Actually, let's get the correct prefix from the map
        # This is fragile, let me use a reverse map
        prefix_lookup = {v: k for k, v in PERM_PREFIX_MAP.items()}
        prefix = prefix_lookup[svc['code']]
        perm_key = f'{prefix}:{action}'

        row_vals = [ACTION_LABELS[action]]
        for rn in role_names_short:
            row_vals.append('✓' if role_has_perm(rn, perm_key) else '—')
        add_alt_row(tbl, row_vals, font_size=7, bold_first=True)

    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════
#  SECTION 5: REVISION HISTORY & NOTES
# ═══════════════════════════════════════════════
add_section_heading('5.  Revision History & Notes', 1)

add_section_heading('Revision History', 2)
rev_headers = ['Version', 'Date', 'Author', 'Changes']
rev_data = [
    ['v2.0', 'May 2026', 'BA Team', 'Consolidated to 6 canonical actions (View, Create, Update, Delete, Export/Download, Print). Removed granular/inconsistent permission keys. Simplified matrix format.'],
    ['v1.0', 'May 2026', 'BA Team', 'Initial draft — 21 modules mapped, 10 roles defined with granular permission keys.'],
]
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(tbl, rev_headers, font_size=8)
for rd in rev_data:
    add_alt_row(tbl, rd, font_size=8)
doc.add_paragraph()

add_section_heading('Notes', 2)
notes = [
    'The six canonical actions (View, Create, Update, Delete, Export/Download, Print) apply consistently across all 21 modules.',
    'Not all modules support all six actions — only relevant actions are listed per module (see Section 2).',
    'The Super Admin role bypasses all permission checks and has unrestricted access.',
    'The "Delete" action is intentionally limited to Manager, Production Manager, and above to prevent data loss.',
    'Mfg Costing (S17) is owner-only (Super Admin). Manager and below have no access to cost data — this is enforced at both the UI and API level.',
    'Godown Staff can record dispatch video (cctv:create) but cannot browse the footage library (cctv:view) — intentional security separation.',
    'Read-Only role has no access to Admin & Settings (S19). Audit Trail is restricted to Manager and above.',
    'Role changes take effect immediately after saving. Active user sessions should refresh permissions on next request.',
    'This document should be reviewed and signed off by stakeholders before development begins.',
]
for n in notes:
    add_bullet(n)

# ─── Save ───
output_path = 'D:\\Wireframe CMS_GMMS\\HoorTex_RBAC_Permission_Matrix.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total service modules: {len(services)}')
print(f'Total permission keys: {total_perms}')
print(f'Total roles: {len(roles)}')
