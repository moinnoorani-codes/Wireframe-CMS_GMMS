from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(2)

# ─── Helper functions ───
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')

def add_data_row(table, values, bold_first=True):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(8.5)
                if bold_first and i == 0:
                    run.bold = True
    return row

def add_section_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HOORTEX CMS — ORDER MANAGEMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Permission Matrix')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Document Type:  BA Deliverable\n'
                    'Module:        Order Management (Sales ERP)\n'
                    'Platforms:     Web + Mobile\n'
                    'Date:          May 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════
add_section_heading('Table of Contents', 1)
toc_items = [
    '1. Permission Legend',
    '2. Screen-by-Screen Permission Matrix',
    '   2.1  W-03  Dashboard (Order Metrics)',
    '   2.2  W-14  Order List',
    '   2.3  W-15  Order Detail + Payments',
    '   2.4  W-16R Create Order — Retail',
    '   2.5  W-16W Create Order — Wholesale',
    '   2.6  W-34  Wholesale Approval Queue',
    '   2.7  W-17  LR Console',
    '   2.8  W-18  LR Detail',
    '   2.9  W-20  Payment Records',
    '   2.10 W-36  Daily Reconciliation',
    '   2.11 W-39  Challan Print Preview',
    '   2.12 M-07  Order Picking List (Mobile)',
    '   2.13 M-08  Order Detail (Mobile)',
    '   2.14 M-09  Picking Execution (Mobile)',
    '   2.15 M-10  Dispatch Confirmation (Mobile)',
    '   2.16 M-11  LR Upload (Mobile)',
    '   2.17 M-12  LR Confirmation (Mobile)',
    '   2.18 M-15R Create Order Retail (Mobile)',
    '   2.19 M-15W Create Order Wholesale (Mobile)',
    '   2.20 M-16  Wholesale Approvals (Mobile)',
    '3. Role-Based Permission Grouping',
    '   3.1  Role Definitions',
    '   3.2  Role vs. Screen Matrix',
    '   3.3  Role vs. Action Matrix',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PERMISSION LEGEND
# ═══════════════════════════════════════════════════════════
add_section_heading('1. Permission Legend', 1)

legend_headers = ['Permission', 'Scope of Access']
legend_data = [
    ['View',   'View screen, data tables, details, metrics, documents; apply filters; search; export data'],
    ['Create', 'Create new orders (retail/wholesale), add line items, record payments, upload LR, register customers'],
    ['Update', 'Edit/approve/reject orders, resend SMS, start picking, confirm dispatch, close day, record payments'],
    ['Delete', 'Cancel/reject/void orders, remove records (currently limited to Reject action)'],
    ['Print',  'Print challan documents, print labels (can be merged with View or kept as separate control)'],
    ['—',      'Navigation-only element; no permission required'],
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, legend_headers)
for row_data in legend_data:
    add_data_row(table, row_data)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 2. SCREEN-BY-SCREEN PERMISSION MATRIX
# ═══════════════════════════════════════════════════════════
add_section_heading('2. Screen-by-Screen Permission Matrix', 1)

# ── Screen data ──
screens = [
    {
        'id': '2.1', 'code': 'W-03', 'name': 'Dashboard — Order Metrics',
        'actions': [
            ('View order metrics (Pending Orders, Today Dispatches, etc.)', 'View'),
            ('Refresh', 'View'),
            ('Export dashboard data', 'View'),
            ('Click order alert / activity item', 'View'),
            ('Click stock / payment alert', 'View'),
        ]
    },
    {
        'id': '2.2', 'code': 'W-14', 'name': 'Order List',
        'actions': [
            ('View Order List screen', 'View'),
            ('+ New Order', 'Create'),
            ('Export order list', 'View'),
            ('Filter tabs (All / Pending Approval / Approved / etc.)', 'View'),
            ('Order Type toggle (All / Retail / Wholesale)', 'View'),
            ('View row → navigate to Order Detail', 'View'),
        ]
    },
    {
        'id': '2.3', 'code': 'W-15', 'name': 'Order Detail + Payments',
        'actions': [
            ('View Order Detail screen', 'View'),
            ('Record Payment (open inline dialog)', 'Create'),
            ('Save Payment Record', 'Create'),
            ('Print Challan', 'Print'),
            ('Resend LR SMS', 'Update'),
            ('View LR document', 'View'),
            ('View CCTV recording', 'View'),
            ('View Payment History', 'View'),
            ('Edit order fields (locked — admin override)', 'Update'),
        ]
    },
    {
        'id': '2.4', 'code': 'W-16R', 'name': 'Create Order — Retail',
        'actions': [
            ('Save as Draft', 'Create'),
            ('Submit + Print Challan', 'Create / Print'),
            ('Search existing customer', 'View'),
            ('Quick Add new customer', 'Create'),
            ('Scan / search items', 'View'),
            ('Add item to order (+ Add)', 'Create'),
            ('Enter qty / discount', 'Create'),
            ('Record payment on creation', 'Create'),
        ]
    },
    {
        'id': '2.5', 'code': 'W-16W', 'name': 'Create Order — Wholesale',
        'actions': [
            ('Save as Draft', 'Create'),
            ('Submit Order + Print Challan', 'Create / Print'),
            ('Search wholesale customer', 'View'),
            ('+ Register New Wholesale Customer', 'Create'),
            ('Scan / search items', 'View'),
            ('Add item to order', 'Create'),
            ('Enter payment terms', 'Create'),
            ('Enter logistics info (Broker, Transport, Instructions)', 'Create'),
        ]
    },
    {
        'id': '2.6', 'code': 'W-34', 'name': 'Wholesale Approval Queue',
        'actions': [
            ('View Approval Queue screen', 'View'),
            ('View metrics cards (Pending, Approved Today, etc.)', 'View'),
            ('View Items (per row)', 'View'),
            ('Approve order (per row)', 'Update'),
            ('Reject order (per row)', 'Update'),
            ('Approve Selected (batch)', 'Update'),
            ('Select checkboxes', 'Update'),
        ]
    },
    {
        'id': '2.7', 'code': 'W-17', 'name': 'LR Console',
        'actions': [
            ('View LR Console screen', 'View'),
            ('View LR document (per row)', 'View'),
            ('View Order Detail (per row)', 'View'),
            ('Resend SMS (per row)', 'Update'),
            ('Resend Selected (batch)', 'Update'),
            ('Export LR data', 'View'),
            ('Search / filter', 'View'),
        ]
    },
    {
        'id': '2.8', 'code': 'W-18', 'name': 'LR Detail',
        'actions': [
            ('View LR Detail screen', 'View'),
            ('Resend SMS', 'Update'),
            ('Download LR', 'View'),
            ('View SMS Delivery Log', 'View'),
        ]
    },
    {
        'id': '2.9', 'code': 'W-20', 'name': 'Payment Records',
        'actions': [
            ('View Payment Records screen', 'View'),
            ('View metrics (Outstanding, Credit Limit, etc.)', 'View'),
            ('Order Detail (per row)', 'View'),
            ('Search / filter', 'View'),
        ]
    },
    {
        'id': '2.10', 'code': 'W-36', 'name': 'Daily Reconciliation',
        'actions': [
            ('View Daily Reconciliation screen', 'View'),
            ('Export Statement', 'View'),
            ('Close Day', 'Update'),
            ('View metrics & tables', 'View'),
            ('Filter transaction log', 'View'),
        ]
    },
    {
        'id': '2.11', 'code': 'W-39', 'name': 'Challan Print Preview',
        'actions': [
            ('View Challan Preview', 'View'),
            ('Print Challan', 'Print'),
        ]
    },
    {
        'id': '2.12', 'code': 'M-07', 'name': 'Order Picking List (Mobile)',
        'actions': [
            ('View Pending / In Progress tabs', 'View'),
            ('View order cards', 'View'),
            ('Start Picking', 'Update'),
        ]
    },
    {
        'id': '2.13', 'code': 'M-08', 'name': 'Order Detail (Mobile)',
        'actions': [
            ('View Order Detail', 'View'),
            ('Save Payment (inline form)', 'Create'),
            ('Print Challan', 'Print'),
            ('Begin Picking', 'Update'),
        ]
    },
    {
        'id': '2.14', 'code': 'M-09', 'name': 'Picking Execution (Mobile)',
        'actions': [
            ('View item list & scan progress', 'View'),
            ('Scan item (confirm picked)', 'Update'),
        ]
    },
    {
        'id': '2.15', 'code': 'M-10', 'name': 'Dispatch Confirmation (Mobile)',
        'actions': [
            ('View Dispatch Summary', 'View'),
            ('Confirm Dispatch', 'Update'),
            ('Confirm Dispatch + Print Challan', 'Update / Print'),
        ]
    },
    {
        'id': '2.16', 'code': 'M-11', 'name': 'LR Upload (Mobile)',
        'actions': [
            ('Search order / scan barcode', 'View'),
            ('Take Photo / Upload PDF', 'Create'),
            ('Upload LR + Auto Send SMS', 'Create / Update'),
            ('Resend SMS', 'Update'),
            ('View SMS delivery status', 'View'),
        ]
    },
    {
        'id': '2.17', 'code': 'M-12', 'name': 'LR Confirmation (Mobile)',
        'actions': [
            ('View confirmation screen', 'View'),
            ('Back to Orders', '—'),
        ]
    },
    {
        'id': '2.18', 'code': 'M-15R', 'name': 'Create Order Retail (Mobile)',
        'actions': [
            ('Submit + Print Challan', 'Create / Print'),
            ('Search / Quick Add customer', 'View / Create'),
            ('Scan / search / add item', 'View / Create'),
            ('Record payment', 'Create'),
        ]
    },
    {
        'id': '2.19', 'code': 'M-15W', 'name': 'Create Order Wholesale (Mobile)',
        'actions': [
            ('Submit + Print Challan', 'Create / Print'),
            ('Search / Register customer', 'View / Create'),
            ('Scan / search / add item', 'View / Create'),
            ('Payment & Logistics entries', 'Create'),
        ]
    },
    {
        'id': '2.20', 'code': 'M-16', 'name': 'Wholesale Approvals (Mobile)',
        'actions': [
            ('View Pending / Approved / Rejected tabs', 'View'),
            ('View Items (per order card)', 'View'),
            ('Approve order', 'Update'),
            ('Reject order', 'Update'),
        ]
    },
]

for s in screens:
    add_section_heading(f'{s["id"]}  {s["code"]} — {s["name"]}', 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(14)
        row.cells[1].width = Cm(5)
    add_header_row(tbl, ['Button / Action', 'Required Permission'])
    for action, perm in s['actions']:
        add_data_row(tbl, [action, perm])
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 3. ROLE-BASED PERMISSION GROUPING
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading('3. Role-Based Permission Grouping', 1)

# 3.1 Role Definitions
add_section_heading('3.1  Role Definitions', 2)

role_headers = ['Role', 'Description', 'Granted Permissions']
role_data = [
    ['Order Viewer',
     'Read-only access. Can browse orders, view details, check metrics, and export reports. No ability to create, modify, or approve anything.',
     'View'],
    ['Order Creator',
     'Can create new retail/wholesale orders, add items, record payments, register customers. Cannot approve or modify existing approved orders.',
     'View + Create + Print'],
    ['Order Processor',
     'Warehouse/operations role. Can pick items, confirm dispatch, upload LR, resend SMS. Responsible for fulfilment.',
     'View + Update + Create (LR)'],
    ['Order Approver',
     'Can view and approve/reject wholesale orders in the approval queue. Typically a supervisor/manager role.',
     'View + Update (approve/reject)'],
    ['Order Admin',
     'Full access across all order screens including editing locked orders, closing day, and all administrative actions.',
     'View + Create + Update + Delete + Print'],
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(tbl, role_headers)
for rd in role_data:
    add_data_row(tbl, rd)
doc.add_paragraph()

# 3.2 Role vs Screen Matrix
add_section_heading('3.2  Role vs. Screen Access Matrix', 2)

screen_list = [
    ('W-03',  'Dashboard'),
    ('W-14',  'Order List'),
    ('W-15',  'Order Detail'),
    ('W-16R', 'Create Order Retail'),
    ('W-16W', 'Create Order Wholesale'),
    ('W-34',  'Approval Queue'),
    ('W-17',  'LR Console'),
    ('W-18',  'LR Detail'),
    ('W-20',  'Payment Records'),
    ('W-36',  'Daily Reconciliation'),
    ('W-39',  'Challan Print Preview'),
    ('M-07',  'Order Picking List'),
    ('M-08',  'Order Detail (Mobile)'),
    ('M-09',  'Picking Execution'),
    ('M-10',  'Dispatch Confirmation'),
    ('M-11',  'LR Upload'),
    ('M-12',  'LR Confirmation'),
    ('M-15R', 'Create Order Retail (M)'),
    ('M-15W', 'Create Order Wholesale (M)'),
    ('M-16',  'Wholesale Approvals (M)'),
]

# Matrix data: rows = screens, columns = roles
# Legend: V=View, C=Create, U=Update, D=Delete, P=Print, —=No Access
matrix = {
    'Order Viewer':    ['V', 'V', 'V', '—',  '—',  '—',  'V', 'V', 'V', 'V', 'V', '—',  '—',  '—',  '—',  '—',  '—',  '—',  '—',  '—'],
    'Order Creator':   ['V', 'V+C', 'V', 'V+C+P', 'V+C+P', '—',  'V', 'V', 'V', 'V', 'V+P', '—',  '—',  '—',  '—',  '—',  '—',  'V+C+P', 'V+C+P', '—'],
    'Order Processor': ['V', 'V', 'V+U', '—',  '—',  '—',  'V+U', 'V+U', 'V', 'V', 'V+P', 'V+U', 'V+U', 'V+U', 'V+U', 'V+C+U', 'V', '—',  '—',  '—'],
    'Order Approver':  ['V', 'V', 'V', '—',  '—',  'V+U', 'V', 'V', 'V', 'V', 'V', '—',  '—',  '—',  '—',  '—',  '—',  '—',  '—',  'V+U'],
    'Order Admin':     ['V', 'V+C', 'V+C+U+D+P', 'V+C+P', 'V+C+P', 'V+U', 'V+U', 'V+U', 'V+C', 'V+U', 'V+P', 'V+U', 'V+C+U+P', 'V+U', 'V+U+P', 'V+C+U', 'V', 'V+C+P', 'V+C+P', 'V+U'],
}

num_cols = 2 + len(matrix)
tbl = doc.add_table(rows=1, cols=num_cols)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(tbl, ['Screen', 'Code'] + list(matrix.keys()))

for i, (code, name) in enumerate(screen_list):
    row_vals = [f'{code} — {name}']
    for role_name in matrix.keys():
        row_vals.append(matrix[role_name][i])
    add_data_row(tbl, row_vals, bold_first=True)

doc.add_paragraph()

# 3.3 Role vs Action Matrix
add_section_heading('3.3  Role vs. Action (CRUD) Summary', 2)

action_headers = ['Action']
action_rows = ['View Screens & Data', 'Create Orders', 'Create Payments', 'Create Customers', 'Create LR Documents',
               'Update / Edit Orders', 'Approve Orders', 'Reject Orders', 'Resend SMS',
               'Start Picking', 'Confirm Dispatch', 'Close Day', 'Delete / Cancel Orders', 'Print Challan']

role_actions = {
    'Order Viewer':    ['✓', '—', '—', '—', '—', '—', '—', '—', '—', '—', '—', '—', '—', '—'],
    'Order Creator':   ['✓', '✓', '✓', '✓', '—', '—', '—', '—', '—', '—', '—', '—', '—', '✓'],
    'Order Processor': ['✓', '—', '—', '—', '✓', '✓', '—', '—', '✓', '✓', '✓', '—', '—', '✓'],
    'Order Approver':  ['✓', '—', '—', '—', '—', '—', '✓', '✓', '—', '—', '—', '—', '—', '—'],
    'Order Admin':     ['✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓'],
}

tbl = doc.add_table(rows=1, cols=6)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(tbl, ['Action'] + list(role_actions.keys()))

for i, action in enumerate(action_rows):
    row_vals = [action]
    for role_name in role_actions.keys():
        row_vals.append(role_actions[role_name][i])
    add_data_row(tbl, row_vals, bold_first=True)

doc.add_paragraph()

# ─── Legend for role matrix ───
legend_p = doc.add_paragraph()
legend_p.add_run('Legend:  ').bold = True
legend_p.add_run('✓ = Granted    — = Not Granted')
legend_p.runs[1].font.size = Pt(9)

doc.add_paragraph()

# ─── Footer note ───
add_section_heading('Notes', 2)
notes = [
    'Permissions are additive — each higher role includes all permissions of lower roles where marked.',
    '"Print" is listed as a separate permission but can be merged into "View" or "Update" per business preference.',
    '"Delete" is intentionally restricted to Admin only. Order rejection is treated as a soft-delete / status change (Update), not a hard delete.',
    'Role assignments should be configurable per user in the Admin Settings module.',
    'This matrix covers ONLY order-related screens. Other modules (Inventory, CRM, Manufacturing, HR) require separate matrices.',
]
for n in notes:
    p = doc.add_paragraph(n, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(9)

# ─── Save ───
output_path = 'D:\\Wireframe CMS_GMMS\\HoorTex_Order_Permission_Matrix.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
